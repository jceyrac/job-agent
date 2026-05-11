#!/usr/bin/env python3
"""
prepare.py — application prep module for job_agent

Generates: cover letter, CV bullet selection, company research notes,
and screening Q&A drafts for jobs in 'ready' status.

Usage:
    python prepare.py --job <job_id>
    python prepare.py --job <job_id> --profile <id>
    python prepare.py --ready
    python prepare.py --ready --limit 5
    python prepare.py --job <id> --redo
    python prepare.py --job <id> --mock
"""

import argparse
import json
import os
import re
import sys
import time
from datetime import datetime, timezone

from dotenv import load_dotenv
load_dotenv()

from profiles import ALL_PROFILES
from scorer import (
    _call_groq_fallback_chain,
    _call_deepseek,
)
from storage import JobStorage

DB_PATH = "data/jobs.db"

# ---------------------------------------------------------------------------
# Model chains for preparation tasks
# ---------------------------------------------------------------------------

PREPARE_HEAVY_MODELS = [
    "llama-3.3-70b-versatile",
    "meta-llama/llama-4-scout-17b-16e-instruct",
]

PREPARE_LIGHT_MODELS = [
    "meta-llama/llama-4-scout-17b-16e-instruct",
    "llama-3.3-70b-versatile",
]

# ---------------------------------------------------------------------------
# CV bullet library — extracted from Jérôme's CV
# ---------------------------------------------------------------------------

CV_BULLETS = [
    "Senior Product Manager with 10+ years of experience in designing and delivering SaaS, fintech, DeFi and AI products following lean and design thinking approaches with agile cross-functional teams.",
    "Created a Cloud Service Center (Cloud Assembly Factory) to provide cloud infrastructure based on Microsoft Azure for digital products at Vaudoise Assurances.",
    "Managed the API Bank product, the bank data aggregation market leader in France at Powens (ex-Budget Insight), leading PSD2 API integration covering 90%+ of the French banking market.",
    "Increased end-user conversion rates by 30% through app-to-app authentication redesign at Powens.",
    "Developed B2B services client segment (accounting, ERP, payroll) by conducting discovery research with clients at Powens.",
    "Managed the API Single Invoice Cover (SIC), the first real-time B2B credit-decisioning API in French factoring at Allianz Trade (Euler Hermes Digital Agency).",
    "Collaborated with Crédit Agricole during the launch of their factoring product Cash in Time, using the SIC API to insure invoices.",
    "Managed GENIAC, a platform taking care of SMEs administrative and compliance tasks (accounting, HR, legal), designed and shipped the MVP securing Series A investment.",
    "Conducted a strategic pivot of the GENIAC product to open it to a new segment of customers through HR and payroll services.",
    "Designed and shipped, as a Business Analyst, Crédit du Nord's new online wire transfer application at Accenture.",
    "Managed a team of 10 developers and a €4.5M budget to deliver a new banking loan platform for two retail banks within Société Générale.",
    "Designed and delivered an orchestrator to automate subscription of Data Finder services at Adeo, reducing support team workload.",
    "Designed efficient onboarding workflows allowing non-tech users to easily share their data at Adeo.",
    "Designed, shipped and launched the BoF career website at Business Of Fashion in London.",
    "Release manager and Java developer of Crédit du Nord's new banking professional loan branch application.",
    "Front-end developer of Banque Populaire's new e-banking application.",
    "Engineering background (Master's Degree in Engineering, EPF 2005) with hands-on coding experience in Python, Java, C#, PHP, SQL, Javascript.",
    "DeFi certification from Duke University (2025), Machine Learning from Stanford (2024), UX Design from Google (2024).",
    "Scrum Master certified (2010).",
    "Languages: English (fluent), French (mother tongue), German (intermediate), Spanish (intermediate).",
]

# ---------------------------------------------------------------------------
# Style anchors — key characteristics from Jérôme's cover letters
# ---------------------------------------------------------------------------

STYLE_ANCHORS = """
## Style anchors for Jérôme's cover letters

Voice characteristics observed across his cover letters:
- Opens with "I'm Jérôme, a Product Manager with more than 10 years of experience..."
- Direct and confident without being arrogant
- Connects specific past experience to the target role's requirements
- Names concrete achievements with metrics when possible (30% conversion lift, 90% market coverage)
- References the company by name and states why THIS company specifically
- Closes with "Looking forward to having a chat soon" or similar
- Uses "Yours sincerely" or "Warm regards" as sign-off
- Avoids: "I am excited to apply", "passionate about", "dynamic team", "results-driven", "hit the ground running"
- Does NOT use: corporate jargon, buzzwords, exaggerated enthusiasm
- Tone: professional but conversational, like a peer talking to a peer
- Length: concise, one page maximum, gets to the point quickly
"""

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

COVER_LETTER_SYSTEM = """You are an expert cover letter writer who crafts tailored, authentic application letters.

{style_anchors}

## Critical rules
- Write in {language_name}. Match the JD's language exactly.
- Target 250–350 words. Be concise — every sentence must earn its place.
- Open directly. No "Dear Sir/Madam" boilerplate if a name isn't available — "Dear Hiring Manager" is acceptable.
- Make at least TWO specific references to concrete responsibilities or requirements from the job description. Show you read it.
- NEVER use these phrases: "I am excited to apply", "passionate about", "dynamic team", "results-driven", "hit the ground running", "think outside the box", "go-getter", "synergy", "game-changer".
- Connect the candidate's actual experience to what the role needs. Be specific about which experience matches which requirement.
- Voice: direct, confident, conversational. Peer to peer, not supplicant to gatekeeper.
- Close with "Warm regards," followed by the candidate's name on the next line.

## Candidate context
{scoring_context}

## CV bullet library (select relevant ones, don't list all)
{bullets}

Return ONLY the cover letter text — no preamble, no "Here is your cover letter", no markdown headers."""


SCREENING_SYSTEM = """You are helping a candidate draft concise, authentic answers to common screening questions.

{style_anchors}

## Critical rules
- Write in {language_name}.
- Each answer: 80–150 words. Tight, specific, no filler.
- NEVER use these phrases: "I am excited to", "passionate about", "dynamic team", "results-driven", "hit the ground running", "think outside the box".
- Ground every answer in the job description and the candidate's actual experience.
- Voice: direct, confident, conversational.

## Candidate context
{scoring_context}

## CV bullet library
{bullets}

Return ONLY a JSON object with these keys (no preamble, no markdown):
{{
  "why_company": "<80-150 word answer to 'Why do you want to work at this company?'>",
  "why_fit": "<80-150 word answer to 'Why are you a good fit for this role?'>",
  "salary": "<80-150 word framing of salary expectations — professional, non-committal, references market rates for the role/seniority/location without giving a specific number unless the JD demands one>",
  "notice_period": "<80-150 word framing of notice period / availability — standard Swiss/EU notice periods, flexible start date language>"
}}"""


BULLET_SYSTEM = """You are an expert CV reviewer matching a candidate's experience bullets to a specific job description.

Given the job description and the candidate's full CV bullet library, return a JSON object selecting which bullets to surface (with rationale) and which to omit (with reason), ordered by relevance to THIS specific job.

## Candidate context
{scoring_context}

## Rules
- Select bullets that directly match requirements, responsibilities, or context in the job description.
- Order selected bullets by relevance to THIS job (most relevant first), NOT chronologically.
- Omit bullets that are irrelevant to this specific role. Give a brief reason.
- Omit bullets that would weaken the application (e.g., highlighting non-PM roles for a senior PM position).
- Be selective — typically 5–10 bullets should be selected, not all 20.

Return ONLY a JSON object (no preamble, no markdown):
{{
  "bullets": [
    {{"text": "<exact bullet text from the library>", "rationale": "<why this matches a specific JD requirement>"}}
  ],
  "omit": [
    {{"text": "<exact bullet text from the library>", "reason": "<why this is not relevant to this specific role>"}}
  ]
}}"""


COMPANY_RESEARCH_SYSTEM = """You are a job application researcher. Given ONLY the job description text, extract and synthesize information about the company and the role.

## Critical rules
- Ground EVERYTHING in the job description text. Do not invent, guess, or use external knowledge.
- If the JD doesn't mention something, say so explicitly — don't fabricate.
- Write in {language_name}.

Return ONLY a JSON object (no preamble, no markdown):
{{
  "what_company_does": "<2-3 sentences from the JD about the company's product, market, or mission>",
  "role_stakes": "<2-3 sentences about what success looks like in this role and why it matters to the company, inferred from the JD>",
  "why_jerome_fit": "<2-3 sentences connecting Jérôme's specific experience (from the context below) to this role's requirements, grounded in what the JD asks for>"
}}

## Candidate context
{scoring_context}

## CV bullet library
{bullets}"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _slugify(text: str) -> str:
    """Aggressive ASCII slug: lowercase, dashes, strip non-alphanumeric."""
    text = text.lower().strip()
    text = re.sub(r'[^a-z0-9\s-]', '', text)
    text = re.sub(r'[\s_]+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')[:80]


def _language_name(code: str) -> str:
    """ISO 2-letter code → full language name for prompting."""
    mapping = {
        "en": "English", "fr": "French", "de": "German",
        "it": "Italian", "es": "Spanish",
    }
    return mapping.get(code, "English")


def _load_cv_bullets() -> str:
    """Load CV bullet library from DOCX if available, otherwise use hardcoded list."""
    try:
        from docx import Document
        docx_paths = [
            os.path.expanduser(
                "~/Nextcloud/Documents/01 Job/00 CV+Motiv/CV/EN/Adresse Suisse/"
                "OLD/CV_Generalist_Jérôme_Ceyrac _EN Suisse v0.1.docx"
            ),
            os.path.expanduser(
                "~/Nextcloud/Documents/01 Job/00 CV+Motiv/CV/EN/Adresse Suisse/"
                "OLD/CV_Web3 _Jérôme_Ceyrac _ENCH v0.1.docx"
            ),
        ]
        for path in docx_paths:
            if os.path.exists(path):
                doc = Document(path)
                lines = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text and len(text) > 40:
                        lines.append(text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text and len(text) > 40 and text not in lines:
                                lines.append(text)
                if lines:
                    return "\n".join(f"- {line}" for line in lines[:50])
    except Exception:
        pass
    return "\n".join(f"- {b}" for b in CV_BULLETS)


def _build_user_prompt(job: dict) -> str:
    """Build the user prompt with job details and extracted fields."""
    parts = [
        f"Title: {job.get('title', '')}",
        f"Company: {job.get('company', '')}",
        f"Location: {job.get('location', '')}",
        f"Base location: {job.get('base_location', '')}",
        f"URL: {job.get('url', '')}",
        "",
        f"Industry sector: {job.get('industry_sector', 'unknown')}",
        f"Work mode: {job.get('work_mode', 'unknown')}",
        f"Geo zone: {job.get('geo_zone', 'unknown')}",
        f"Company size: {job.get('company_size', 'unknown')}",
        f"Contract type: {job.get('contract_type', 'unknown')}",
        f"Company country: {job.get('company_country', 'unknown')}",
        f"Language required: {job.get('language_required', 'unknown')}",
        "",
        f"Summary: {job.get('summary', '')}",
        "",
        "Full job description:",
        job.get('description', '') or '',
    ]
    return "\n".join(parts)


def _call_prepare_model(system: str, user: str,
                        preferred_models: list[str],
                        json_mode: bool = True,
                        max_tokens: int = 800) -> tuple[str, str]:
    """
    Call LLM with Groq → DeepSeek fallback chain for preparation tasks.

    Returns (response_text, model_name).
    """
    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]

    # Try Groq chain first
    try:
        raw, model = _call_groq_fallback_chain(
            messages, models=preferred_models,
            json_mode=json_mode, max_tokens=max_tokens,
        )
        return raw, model
    except Exception as e:
        if "All Groq models exhausted" not in str(e):
            raise

    # Fall back to DeepSeek
    print(f"  ⚠️  Groq exhausted — falling back to DeepSeek")
    raw = _call_deepseek(messages, model="deepseek-v4-pro",
                         json_mode=json_mode, max_tokens=max_tokens)
    time.sleep(1)
    return raw, "tier_fallback:deepseek-v4-pro"


def _safe_json_parse(raw: str, label: str) -> dict | None:
    """Parse JSON from LLM response, with one retry on failure."""
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Try extracting JSON from markdown code blocks
    m = re.search(r'```(?:json)?\s*([\s\S]*?)```', raw)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass

    # Try finding a JSON object with balanced braces
    brace_start = raw.find('{')
    if brace_start >= 0:
        depth = 0
        for i in range(brace_start, len(raw)):
            if raw[i] == '{':
                depth += 1
            elif raw[i] == '}':
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(raw[brace_start:i + 1])
                    except json.JSONDecodeError:
                        break

    print(f"  ⚠️  Could not parse {label} JSON from: {raw[:200]}...")
    return None


# ---------------------------------------------------------------------------
# Main preparation logic
# ---------------------------------------------------------------------------

def prepare_job_application(job_id: str, profile_id: str | None = None,
                            redo: bool = False, mock: bool = False) -> dict | None:
    """
    Generate the full application package for a job.

    Returns a dict with all generated content, or None if the job doesn't exist
    or is already prepared (without --redo).
    """
    db = JobStorage(DB_PATH)

    job = db.get_job_for_prepare(job_id)
    if job is None:
        print(f"Job '{job_id}' not found in database.")
        return None

    # Auto-pick profile from highest score if not specified
    if profile_id is None:
        profile_id = _auto_pick_profile(db, job_id)
        if profile_id is None:
            print(f"No scored profile found for job '{job_id}'. Specify --profile.")
            return None

    if profile_id not in ALL_PROFILES:
        print(f"Unknown profile '{profile_id}'. Valid: {list(ALL_PROFILES.keys())}")
        return None

    profile = ALL_PROFILES[profile_id]

    # Idempotency check
    existing = db.get_application(job_id)
    if existing and existing.get("prepared_at") and not redo:
        print(f"Job '{job_id}' already prepared (--redo to overwrite).")
        return None

    # Determine language
    language = (job.get("language_required") or "en").strip().lower()
    if language not in ("en", "fr", "de", "it", "es"):
        language = "en"
    language_name = _language_name(language)

    # Load context
    scoring_context = profile.scoring_context or ""
    bullets = _load_cv_bullets()
    user_prompt = _build_user_prompt(job)

    style_context = STYLE_ANCHORS

    print(f"\n{'[MOCK] ' if mock else ''}Preparing: {job.get('title', '')[:60]} @ {job.get('company', '')[:40]}")
    print(f"  Profile: {profile.name} ({profile.id})")
    print(f"  Language: {language_name}")

    result = {
        "job_id": job_id,
        "profile_id": profile_id,
        "language": language,
        "models_used": {},
        "cover_letter": "",
        "cv_bullets_selected": None,
        "company_research": "",
        "screening_answers": None,
        "prepared_by": "",
    }

    # ── 1. Cover letter ────────────────────────────────────────────────────
    print("  Generating cover letter...")
    cover_system = COVER_LETTER_SYSTEM.format(
        style_anchors=style_context,
        language_name=language_name,
        scoring_context=scoring_context,
        bullets=bullets,
    )
    try:
        raw, model = _call_prepare_model(
            cover_system, user_prompt,
            preferred_models=PREPARE_HEAVY_MODELS,
            json_mode=False,
            max_tokens=1500,
        )
        result["cover_letter"] = raw.strip()
        result["models_used"]["cover_letter"] = model
        print(f"    ✅ Cover letter ({model}) — {len(raw.split())} words")
    except Exception as e:
        print(f"    ❌ Cover letter failed: {e}")
        result["cover_letter"] = f"[Generation failed: {e}]"

    # ── 2. Screening answers ────────────────────────────────────────────────
    print("  Generating screening answers...")
    screen_system = SCREENING_SYSTEM.format(
        style_anchors=style_context,
        language_name=language_name,
        scoring_context=scoring_context,
        bullets=bullets,
    )
    try:
        raw, model = _call_prepare_model(
            screen_system, user_prompt,
            preferred_models=PREPARE_HEAVY_MODELS,
            json_mode=True,
            max_tokens=1000,
        )
        parsed = _safe_json_parse(raw, "screening")
        if parsed is None:
            # Retry once with explicit JSON instruction
            retry_system = screen_system + "\n\nIMPORTANT: You MUST respond with ONLY valid JSON. No markdown, no explanation, just the JSON object."
            raw2, model2 = _call_prepare_model(
                retry_system, user_prompt,
                preferred_models=PREPARE_LIGHT_MODELS,
                json_mode=True,
                max_tokens=1000,
            )
            parsed = _safe_json_parse(raw2, "screening-retry")
            if parsed is None:
                parsed = {"why_company": raw, "why_fit": "", "salary": "", "notice_period": ""}
            else:
                model = model2
        result["screening_answers"] = parsed
        result["models_used"]["screening"] = model
        print(f"    ✅ Screening answers ({model})")
    except Exception as e:
        print(f"    ❌ Screening answers failed: {e}")
        result["screening_answers"] = {"error": str(e)}

    # ── 3. CV bullet selection ──────────────────────────────────────────────
    print("  Selecting CV bullets...")
    bullet_system = BULLET_SYSTEM.format(
        style_anchors=style_context,
        scoring_context=scoring_context,
    )
    bullet_prompt = f"{user_prompt}\n\n## CV Bullet Library\n{bullets}"
    try:
        raw, model = _call_prepare_model(
            bullet_system, bullet_prompt,
            preferred_models=PREPARE_LIGHT_MODELS,
            json_mode=True,
            max_tokens=1500,
        )
        parsed = _safe_json_parse(raw, "bullets")
        if parsed is None:
            retry_system = bullet_system + "\n\nIMPORTANT: You MUST respond with ONLY valid JSON. No markdown, no explanation, just the JSON object."
            raw2, model2 = _call_prepare_model(
                retry_system, bullet_prompt,
                preferred_models=PREPARE_LIGHT_MODELS,
                json_mode=True,
                max_tokens=1500,
            )
            parsed = _safe_json_parse(raw2, "bullets-retry")
            if parsed is None:
                parsed = {"bullets": [], "omit": [], "_parse_error": True}
            else:
                model = model2
        result["cv_bullets_selected"] = parsed
        result["models_used"]["bullets"] = model
        n_selected = len(parsed.get("bullets", []))
        n_omitted = len(parsed.get("omit", []))
        print(f"    ✅ Bullet selection ({model}) — {n_selected} selected, {n_omitted} omitted")
    except Exception as e:
        print(f"    ❌ Bullet selection failed: {e}")
        result["cv_bullets_selected"] = {"bullets": [], "omit": [], "error": str(e)}

    # ── 4. Company research ─────────────────────────────────────────────────
    print("  Generating company research...")
    research_system = COMPANY_RESEARCH_SYSTEM.format(
        language_name=language_name,
        scoring_context=scoring_context,
        bullets=bullets,
    )
    try:
        raw, model = _call_prepare_model(
            research_system, user_prompt,
            preferred_models=PREPARE_LIGHT_MODELS,
            json_mode=True,
            max_tokens=600,
        )
        parsed = _safe_json_parse(raw, "research")
        if parsed is None:
            retry_system = research_system + "\n\nIMPORTANT: You MUST respond with ONLY valid JSON."
            raw2, model2 = _call_prepare_model(
                retry_system, user_prompt,
                preferred_models=PREPARE_LIGHT_MODELS,
                json_mode=True,
                max_tokens=600,
            )
            parsed = _safe_json_parse(raw2, "research-retry")
            if parsed is None:
                parsed = {"what_company_does": raw, "role_stakes": "", "why_jerome_fit": ""}
            else:
                model = model2
        result["company_research"] = parsed
        result["models_used"]["research"] = model
        print(f"    ✅ Company research ({model})")
    except Exception as e:
        print(f"    ❌ Company research failed: {e}")
        result["company_research"] = {"error": str(e)}

    # Build the prepared_by summary
    models_used = result["models_used"]
    fallback_count = sum(1 for m in models_used.values() if "fallback" in m)
    result["prepared_by"] = ", ".join(
        f"{task}:{model}" for task, model in sorted(models_used.items())
    )
    result["fallback_count"] = fallback_count

    # ── Persist ─────────────────────────────────────────────────────────────
    if not mock:
        db.save_prepared_application(
            job_id=job_id,
            profile_id=profile_id,
            cover_letter=result["cover_letter"],
            cv_bullets_selected=result["cv_bullets_selected"],
            company_research=result["company_research"],
            screening_answers=result["screening_answers"],
            language=language,
            prepared_by=result["prepared_by"],
        )
        _write_markdown(job, result)
        print(f"  💾 Saved to DB + outputs/applications/")
    else:
        print(f"\n  [MOCK] Cover letter preview:")
        print(f"  {'─' * 60}")
        for line in result["cover_letter"].split("\n")[:15]:
            print(f"  {line}")
        if len(result["cover_letter"].split("\n")) > 15:
            print(f"  ... ({len(result['cover_letter'].split())} words total)")
        print(f"  {'─' * 60}")

    return result


def _auto_pick_profile(db: JobStorage, job_id: str) -> str | None:
    """Pick the profile with the highest score for this job."""
    with db._conn() as conn:
        row = conn.execute(
            """SELECT profile_id FROM job_scores
               WHERE job_id = ? AND score IS NOT NULL
               ORDER BY score DESC LIMIT 1""",
            (job_id,),
        ).fetchone()
    if row:
        pid = row["profile_id"]
        if pid in ALL_PROFILES:
            return pid
    return None


def _write_markdown(job: dict, result: dict) -> None:
    """Write the application package to a markdown file."""
    output_dir = os.path.join(os.path.dirname(__file__), "outputs", "applications")
    os.makedirs(output_dir, exist_ok=True)

    company_slug = _slugify(job.get("company", "unknown"))
    title_slug = _slugify(job.get("title", "untitled"))
    filename = f"{job.get('id', 'unknown')}__{company_slug}__{title_slug}.md"
    filepath = os.path.join(output_dir, filename)

    lang_name = _language_name(result.get("language", "en"))

    lines = [
        f"# Application: {job.get('title', '')} @ {job.get('company', '')}",
        "",
        f"- **URL**: {job.get('url', '')}",
        f"- **Company**: {job.get('company', '')}",
        f"- **Title**: {job.get('title', '')}",
        f"- **Language**: {lang_name}",
        f"- **Model**: {result.get('prepared_by', '')}",
        f"- **Generated**: {datetime.now(timezone.utc).isoformat()}",
        "",
        "---",
        "",
        "## Cover Letter",
        "",
        result.get("cover_letter", ""),
        "",
        "---",
        "",
        "## Screening Answers",
        "",
    ]

    screening = result.get("screening_answers") or {}
    labels = {
        "why_company": "Why this company?",
        "why_fit": "Why are you a fit?",
        "salary": "Salary expectations",
        "notice_period": "Notice period",
    }
    for key, label in labels.items():
        val = screening.get(key, "")
        if isinstance(val, str) and val.strip():
            lines.append(f"### {label}")
            lines.append("")
            lines.append(val)
            lines.append("")

    lines.extend([
        "---",
        "",
        "## CV Bullets",
        "",
    ])

    bullets_data = result.get("cv_bullets_selected") or {}
    selected = bullets_data.get("bullets", [])
    omitted = bullets_data.get("omit", [])

    if selected:
        lines.append("### Selected (in priority order)")
        lines.append("")
        for i, b in enumerate(selected, 1):
            text = b.get("text", "")
            rationale = b.get("rationale", "")
            lines.append(f"{i}. **{text}**")
            if rationale:
                lines.append(f"   - *{rationale}*")
            lines.append("")

    if omitted:
        lines.append("### Omitted")
        lines.append("")
        for b in omitted:
            text = b.get("text", "")
            reason = b.get("reason", "")
            lines.append(f"- ~~{text}~~")
            if reason:
                lines.append(f"  - *{reason}*")
        lines.append("")

    lines.extend([
        "---",
        "",
        "## Company Research",
        "",
    ])

    research = result.get("company_research") or {}
    if isinstance(research, dict):
        for key, label in [
            ("what_company_does", "What the company does"),
            ("role_stakes", "Role stakes"),
            ("why_jerome_fit", "Why Jérôme is a fit"),
        ]:
            val = research.get(key, "")
            if isinstance(val, str) and val.strip():
                lines.append(f"### {label}")
                lines.append("")
                lines.append(val)
                lines.append("")
    elif isinstance(research, str):
        lines.append(research)
        lines.append("")

    with open(filepath, "w") as f:
        f.write("\n".join(lines))

    print(f"  📄 {filepath}")


# ---------------------------------------------------------------------------
# Batch mode: prepare all 'ready' jobs
# ---------------------------------------------------------------------------

def _run_ready(limit: int | None = None) -> None:
    """Prepare all jobs with status='ready' that haven't been prepared yet."""
    db = JobStorage(DB_PATH)
    jobs = db.get_ready_jobs_unprepared()

    if not jobs:
        print("No unprepared 'ready' jobs found.")
        return

    print(f"Found {len(jobs)} unprepared 'ready' jobs")

    if limit and len(jobs) > limit:
        print(f"  --limit {limit}: processing {limit}/{len(jobs)} "
              f"({len(jobs) - limit} deferred to next run)")
        jobs = jobs[:limit]

    prepared = 0
    errors = 0
    models_used: dict[str, int] = {}
    fallback_total = 0
    languages: dict[str, int] = {}

    for i, job in enumerate(jobs, 1):
        print(f"\n[{i}/{len(jobs)}] {job.get('title', '')[:60]} @ {job.get('company', '')[:40]}")

        try:
            result = prepare_job_application(
                job_id=job["id"],
                profile_id=None,
                redo=False,
                mock=False,
            )
            if result:
                prepared += 1
                fallback_total += result.get("fallback_count", 0)
                lang = result.get("language", "en")
                languages[lang] = languages.get(lang, 0) + 1
                for model in result.get("models_used", {}).values():
                    models_used[model] = models_used.get(model, 0) + 1
            else:
                errors += 1
        except Exception as e:
            print(f"  ❌ Failed: {e}")
            errors += 1

        if i < len(jobs):
            time.sleep(4)

    print(f"\n{'=' * 60}")
    print(f"Preparation complete: {prepared} prepared, {errors} errors")
    if models_used:
        print(f"Models used: {dict(models_used)}")
    print(f"Fallback calls (DeepSeek): {fallback_total}")
    print(f"Languages: {dict(languages)}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="prepare.py — generate application packages for jobs"
    )
    parser.add_argument("--job", default=None, help="Job ID to prepare")
    parser.add_argument("--profile", default=None, help="Profile ID to use")
    parser.add_argument("--ready", action="store_true",
                        help="Prepare all jobs with status='ready'")
    parser.add_argument("--limit", type=int, default=None, metavar="N",
                        help="Cap number of jobs processed")
    parser.add_argument("--redo", action="store_true",
                        help="Overwrite existing application")
    parser.add_argument("--mock", action="store_true",
                        help="Dry-run: print to stdout, no DB or file writes")
    args = parser.parse_args()

    if not args.job and not args.ready:
        print("Specify --job <id> or --ready.")
        sys.exit(1)

    if args.job and args.ready:
        print("--job and --ready are mutually exclusive.")
        sys.exit(1)

    if args.mock and args.ready:
        print("--mock is only supported with --job.")
        sys.exit(1)

    if args.profile and args.profile not in ALL_PROFILES:
        print(f"Unknown profile '{args.profile}'. Valid: {list(ALL_PROFILES.keys())}")
        sys.exit(1)

    if not os.path.exists(DB_PATH):
        print("DB not found — run scrape.py first.")
        sys.exit(1)

    if args.ready:
        _run_ready(args.limit)
        return

    # Single job mode
    result = prepare_job_application(
        job_id=args.job,
        profile_id=args.profile,
        redo=args.redo,
        mock=args.mock,
    )

    if result:
        print(f"\nDone. Prepared by: {result['prepared_by']}")


if __name__ == "__main__":
    main()
